import os
import json
from pathlib import Path
import questionary
import tkinter as tk
from tkinter import filedialog
import whisper
import re
import subprocess
from typing import List

try:
    from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
except ImportError:
    AutoTokenizer = None
    AutoModelForSeq2SeqLM = None
    pipeline = None

WHISPER_MODELS = [
    "tiny", "base", "small", "medium", "large", "turbo", "large-v2", "large-v3"
]
CONFIG_FILE = os.path.join(str(Path.home()), ".whisper_config.json")
WHISPER_FOLDER_NAME = "whisper"


def get_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    return {}

def save_config(config):
    with open(CONFIG_FILE, "w") as f:
        json.dump(config, f)

def select_model_folder():
    root = tk.Tk()
    root.withdraw()
    folder = filedialog.askdirectory(title="Select folder to store Whisper models")
    root.destroy()
    return folder

def get_model_folder():
    config = get_config()
    if "model_dir" in config and os.path.exists(config["model_dir"]):
        return config["model_dir"]
    folder = select_model_folder()
    if not folder:
        raise Exception("No folder selected.")
    whisper_dir = os.path.join(folder, WHISPER_FOLDER_NAME)
    os.makedirs(whisper_dir, exist_ok=True)
    config["model_dir"] = whisper_dir
    save_config(config)
    return whisper_dir

def is_model_downloaded(model_name, model_dir):
    # This is a placeholder check; adjust as needed for your model format
    model_path = os.path.join(model_dir, f"{model_name}.pt")
    return os.path.exists(model_path)

def download_model(model_name, model_dir):
    print(f"Downloading model '{model_name}' to {model_dir}...")
    model = whisper.load_model(model_name, download_root=model_dir)
    print("Download complete.")
    return model

def select_model():
    try:
        import sys
        if sys.stdin.isatty():
            return questionary.select(
                "Which Whisper model do you want to use?",
                choices=WHISPER_MODELS
            ).ask()
        else:
            raise Exception("No TTY")
    except Exception:
        print("No interactive terminal detected. Please type the model name from the following list:")
        print(", ".join(WHISPER_MODELS))
        while True:
            model = input("Model name: ").strip()
            if model in WHISPER_MODELS:
                return model
            print("Invalid model. Try again.")

def get_whisper_model():
    model_dir = get_model_folder()
    model_name = select_model()
    model_path = os.path.join(model_dir, f"{model_name}.pt")
    if not is_model_downloaded(model_name, model_dir):
        model = download_model(model_name, model_dir)
    else:
        model = whisper.load_model(model_name, download_root=model_dir)
    import torch
    if torch.cuda.is_available():
        model = model.to('cuda')
        print("Using CUDA (GPU) for Whisper model.")
    else:
        print("CUDA not available, using CPU for Whisper model.")
    print(f"Model '{model_name}' is ready at {model_dir}")
    return model

# Task queue for batch processing
_tasks = []

def is_url(s):
    # Dynamically check if the string is a URL using urllib.parse
        from urllib.parse import urlparse
        if not isinstance(s, str):
            return False
        parsed = urlparse(s)
        return bool(parsed.scheme and parsed.netloc) or re.match(r"^[a-zA-Z0-9\-]+\.[a-zA-Z]{2,}", s)

def download_from_url(url, out_dir="downloads"):
    """
    Download audio (mp3) and audio+video (mp4) from a URL (video or playlist) using yt-dlp.
    Returns list of downloaded mp3 file paths.
    """
    os.makedirs(out_dir, exist_ok=True)
    # Check if playlist
    is_playlist = False
    try:
        # yt-dlp --flat-playlist -J <url> will output JSON with _type=playlist if playlist
        result = subprocess.run([
            "yt-dlp", "--flat-playlist", "-J", url
        ], capture_output=True, text=True, check=True)
        import json as _json
        info = _json.loads(result.stdout)
        is_playlist = info.get("_type") == "playlist"
    except Exception:
        pass
    # Download mp3 (audio only)
    mp3_cmd = [
        "yt-dlp", "-f", "bestaudio", "--extract-audio", "--audio-format", "mp3",
        "-o", os.path.join(out_dir, "%(title)s.%(ext)s"), url
    ]
    subprocess.run(mp3_cmd, check=True)
    # Download mp4 (audio+video)
    mp4_cmd = [
        "yt-dlp", "-f", "bestvideo+bestaudio/best", "-o",
        os.path.join(out_dir, "%(title)s.%(ext)s"), url
    ]
    subprocess.run(mp4_cmd, check=True)
    # Collect all mp3 files in out_dir
    mp3_files = [os.path.join(out_dir, f) for f in os.listdir(out_dir) if f.endswith('.mp3')]
    return mp3_files

def add_task(file, job, tran=None):
    """
    Add a transcription task to the queue.
    :param file: Path to the audio file to transcribe or a URL.
    :param job: Output format(s), comma-separated.
    :param tran: Translation target language(s), comma-separated (e.g. 'fra_Latn,spa_Latn').
    """
    if is_url(file):
        print(f"Detected URL: {file}. Downloading audio and video...")
        mp3_files = download_from_url(file)
        for mp3 in mp3_files:
            _tasks.append({'file': mp3, 'job': job, 'tran': tran})
    else:
        _tasks.append({'file': file, 'job': job, 'tran': tran})


_translation_pipeline = None
_translation_tokenizer = None
_translation_model = None

def get_translation_pipeline():
    global _translation_pipeline, _translation_tokenizer, _translation_model
    if _translation_pipeline is not None:
        return _translation_pipeline
    if AutoTokenizer is None or AutoModelForSeq2SeqLM is None or pipeline is None:
        raise ImportError("transformers not installed. Please install with 'pip install transformers sentencepiece'.")
    model_name = "facebook/nllb-200-distilled-600M"
    # Use the same directory as Whisper models
    cache_dir = get_model_folder()
    _translation_tokenizer = AutoTokenizer.from_pretrained(model_name, cache_dir=cache_dir)
    _translation_model = AutoModelForSeq2SeqLM.from_pretrained(model_name, cache_dir=cache_dir)
    # Use GPU if available
    import torch
    device = 0 if torch.cuda.is_available() else -1
    _translation_pipeline = pipeline("translation", model=_translation_model, tokenizer=_translation_tokenizer, src_lang="eng_Latn", device=device)
    if device == 0:
        print("[INFO] Translation pipeline is using CUDA (GPU).")
    else:
        print("[INFO] Translation pipeline is using CPU.")
    return _translation_pipeline

def split_text_for_translation(text, max_length=1024):
    """Split text into chunks that fit the model's max token length."""
    import re
    # Split by sentences (simple heuristic)
    sentences = re.split(r'(?<=[.!?]) +', text)
    chunks = []
    current = ''
    for s in sentences:
        if len(current) + len(s) < max_length:
            current += (s + ' ')
        else:
            chunks.append(current.strip())
            current = s + ' '
    if current.strip():
        chunks.append(current.strip())
    return chunks

def translate_text(text: str, tgt_langs: List[str]) -> dict:
    """Translate text to multiple target languages using NLLB-200, splitting if too long."""
    pipe = get_translation_pipeline()
    translations = {}
    for lang in tgt_langs:
        try:
            chunks = split_text_for_translation(text)
            translated_chunks = []
            for chunk in chunks:
                result = pipe(chunk, tgt_lang=lang)
                translated_chunks.append(result[0]['translation_text'])
            translations[lang] = ' '.join(translated_chunks)
        except Exception as e:
            translations[lang] = f"[Translation error: {e}]"
    return translations

def translate_segments(segments, tgt_langs):
    """Translate each segment's text for subtitle formats, batching for GPU efficiency, with progress bar."""
    pipe = get_translation_pipeline()
    translations = {lang: [] for lang in tgt_langs}
    try:
        from tqdm import tqdm
        use_tqdm = True
    except ImportError:
        use_tqdm = False
    for lang in tgt_langs:
        texts = [seg.get('text', '') if seg.get('text', '').strip() else '' for seg in segments]
        non_empty_indices = [i for i, t in enumerate(texts) if t.strip()]
        non_empty_texts = [texts[i] for i in non_empty_indices]
        try:
            if non_empty_texts:
                if use_tqdm:
                    print(f"Translating {len(non_empty_texts)} segments to {lang}...")
                # tqdm for visual progress if many segments
                results = []
                batch_size = 32  # reasonable batch size for GPU
                iterator = range(0, len(non_empty_texts), batch_size)
                if use_tqdm:
                    iterator = tqdm(iterator, desc=f"{lang}", unit="batch")
                for i in iterator:
                    batch = non_empty_texts[i:i+batch_size]
                    batch_results = pipe(batch, tgt_lang=lang)
                    results.extend(batch_results)
                translated = [''] * len(texts)
                for idx, seg_idx in enumerate(non_empty_indices):
                    translated[seg_idx] = results[idx]['translation_text']
                translations[lang] = translated
            else:
                translations[lang] = [''] * len(texts)
        except Exception as e:
            translations[lang] = [f"[Translation error: {e}]"] * len(texts)
    return translations

def execute_task():
    """
    Execute all queued transcription tasks using a single Whisper model instance.
    Returns a list of results (dicts with 'file' and 'text').
    Handles output formatting based on job type.
    """
    try:
        from tqdm import tqdm
        use_tqdm = True
    except ImportError:
        use_tqdm = False
    if not _tasks:
        print("No tasks to execute.")
        return []
    model = get_whisper_model()
    results = []
    from collections import defaultdict
    file_jobs = defaultdict(set)
    file_trans = defaultdict(set)
    for task in _tasks:
        jobs = [j.strip() for j in str(task.get('job', '')).split(',') if j.strip()]
        trans = [t.strip() for t in str(task.get('tran', '')).split(',') if t and t.strip()]
        for job in jobs:
            file_jobs[task['file']].add(job)
        for t in trans:
            file_trans[task['file']].add(t)
    file_paths = list(file_jobs.keys())
    file_iter = file_paths
    if use_tqdm:
        file_iter = tqdm(file_paths, desc="Files", unit="file")
    for file_path in file_iter:
        jobs = file_jobs[file_path]
        trans = set(t for t in file_trans.get(file_path, set()) if t and str(t).lower() != "none")
        print(f"Processing: {file_path} (jobs: {', '.join(jobs)})" + (f", translations: {', '.join(trans)}" if trans else ""))
        result = model.transcribe(file_path)
        base = os.path.splitext(os.path.basename(file_path))[0]
        segments = result.get('segments', None)
        out_dir = os.path.join(os.path.dirname(file_path), base)
        os.makedirs(out_dir, exist_ok=True)
        translations = None
        segment_translations = None
        if trans:
            if segments:
                segment_translations = translate_segments(segments, list(trans))
            translations = translate_text(result['text'], list(trans))
        def output_jobs(jobs, base, text, segments, lang_suffix="", translated_segments=None):
            if not text or text is None:
                return
            if use_tqdm:
                job_iter = tqdm(list(jobs), desc=f"Jobs for {base}{lang_suffix}", unit="job")
            else:
                job_iter = jobs
            for job in job_iter:
                print(f"[DEBUG] Processing job: {job} for base: {base}{lang_suffix}")
                # All outputs go into the per-file folder
                if job == 'srt' and translated_segments is not None:
                    try:
                        out_path = os.path.join(out_dir, f"{base}{lang_suffix}.srt")
                        segs = segments if segments else [{'start': 0, 'end': 0, 'text': text}]
                        print(f"[DEBUG] Writing SRT to {out_path} with {len(segs)} segments.")
                        def srt_time(t):
                            h = int(t // 3600)
                            m = int((t % 3600) // 60)
                            s = int(t % 60)
                            ms = int((t - int(t)) * 1000)
                            return f"{h:02}:{m:02}:{s:02},{ms:03}"
                        with open(out_path, 'w', encoding='utf-8') as f:
                            for idx, seg in enumerate(segs, 1):
                                start = seg.get('start', 0)
                                end = seg.get('end', 0)
                                # Use translated segment text
                                seg_text = translated_segments[idx-1] if idx-1 < len(translated_segments) else ''
                                f.write(f"{idx}\n{srt_time(start)} --> {srt_time(end)}\n{seg_text.strip()}\n\n")
                        print(f"[LOG] SRT file successfully saved at: {out_path}")
                    except Exception as e:
                        print(f"[ERROR] Exception while saving SRT: {e}")
                elif job == 'txt':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}_all_transcripts.txt")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(text.strip().replace('\n', ' ') + '\n')
                    print(f"Saved transcript to {out_path}")
                elif job in ('txt with file name', 'twfn'):
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.txt")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(text.strip())
                    print(f"Saved transcript to {out_path}")
                elif job == 'vtt':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.vtt")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write("WEBVTT\n\n")
                        for seg in segments:
                            def vtt_time(t):
                                h = int(t // 3600)
                                m = int((t % 3600) // 60)
                                s = int(t % 60)
                                ms = int((t - int(t)) * 1000)
                                return f"{h:02}:{m:02}:{s:02}.{ms:03}"
                            f.write(f"{vtt_time(seg['start'])} --> {vtt_time(seg['end'])}\n{seg['text'].strip()}\n\n")
                    print(f"Saved VTT to {out_path}")
                elif job == 'json':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.json")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        json.dump(result, f, ensure_ascii=False, indent=2)
                    print(f"Saved JSON to {out_path}")
                elif job == 'csv':
                    import csv
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.csv")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8', newline='') as f:
                        writer = csv.writer(f)
                        writer.writerow(['Index', 'Start', 'End', 'Text'])
                        for idx, seg in enumerate(segments, 1):
                            writer.writerow([idx, seg['start'], seg['end'], seg['text']])
                    print(f"Saved CSV to {out_path}")
                elif job == 'html':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.html")
                    html_content = f"<html><head><meta charset='utf-8'><title>{base}</title></head><body><pre>{result['text'].strip()}</pre></body></html>"
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(html_content)
                    print(f"Saved HTML to {out_path}")
                elif job == 'md':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.md")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(f"# Transcript for {base}\n\n")
                        f.write(result['text'].strip())
                    print(f"Saved Markdown to {out_path}")
                elif job == 'lines':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}_lines.txt")
                    segments = result.get('segments', [{'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        for seg in segments:
                            f.write(seg['text'].strip() + '\n')
                    print(f"Saved line-separated transcript to {out_path}")
                elif job == 'docx':
                    try:
                        from docx import Document
                    except ImportError:
                        print("python-docx not installed. Please install with 'pip install python-docx'.")
                        continue
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.docx")
                    doc = Document()
                    doc.add_heading(f"Transcript for {base}", 0)
                    doc.add_paragraph(result['text'].strip())
                    doc.save(out_path)
                    print(f"Saved DOCX to {out_path}")
                elif job == 'pdf':
                    try:
                        from fpdf import FPDF
                    except ImportError:
                        print("fpdf not installed. Please install with 'pip install fpdf'.")
                        continue
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.pdf")
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    for line in result['text'].splitlines():
                        pdf.cell(0, 10, txt=line, ln=1)
                    pdf.output(out_path)
                    print(f"Saved PDF to {out_path}")
                elif job == 'latex':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.tex")
                    latex = f"""\\documentclass{{article}}\n\\begin{{document}}\n\\section*{{Transcript for {base}}}\n{result['text'].strip()}\n\\end{{document}}\n"""
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(latex)
                    print(f"Saved LaTeX to {out_path}")
                elif job == 'tsv':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.tsv")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8', newline='') as f:
                        f.write('Index\tStart\tEnd\tText\n')
                        for idx, seg in enumerate(segments, 1):
                            f.write(f"{idx}\t{seg['start']}\t{seg['end']}\t{seg['text']}\n")
                    print(f"Saved TSV to {out_path}")
                elif job == 'xml':
                    import xml.etree.ElementTree as ET
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.xml")
                    root = ET.Element('transcript')
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    for idx, seg in enumerate(segments, 1):
                        seg_el = ET.SubElement(root, 'segment', index=str(idx), start=str(seg['start']), end=str(seg['end']))
                        seg_el.text = seg['text']
                    tree = ET.ElementTree(root)
                    tree.write(out_path, encoding='utf-8', xml_declaration=True)
                    print(f"Saved XML to {out_path}")
                elif job == 'yaml':
                    try:
                        import yaml
                    except ImportError:
                        print("pyyaml not installed. Please install with 'pip install pyyaml'.")
                        continue
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.yaml")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        yaml.dump(result, f, allow_unicode=True)
                    print(f"Saved YAML to {out_path}")
                elif job == 'ass':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.ass")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write("[Script Info]\nScriptType: v4.00+\n\n[V4+ Styles]\n\n[Events]\n")
                        for seg in segments:
                            def ass_time(t):
                                h = int(t // 3600)
                                m = int((t % 3600) // 60)
                                s = int(t % 60)
                                cs = int((t - int(t)) * 100)
                                return f"{h:d}:{m:02d}:{s:02d}.{cs:02d}"
                            f.write(f"Dialogue: 0,{ass_time(seg['start'])},{ass_time(seg['end'])},Default,,0,0,0,,{seg['text'].strip()}\n")
                    print(f"Saved ASS to {out_path}")
                elif job == 'sbv':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.sbv")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        for seg in segments:
                            def sbv_time(t):
                                h = int(t // 3600)
                                m = int((t % 3600) // 60)
                                s = int(t % 60)
                                ms = int((t - int(t)) * 100)
                                return f"{h:01d}:{m:02d}:{s:02d}.{ms:02d}"
                            f.write(f"{sbv_time(seg['start'])},{sbv_time(seg['end'])}\n{seg['text'].strip()}\n\n")
                    print(f"Saved SBV to {out_path}")
                elif job in ('dfxp', 'ttml'):
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.dfxp")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write('<?xml version="1.0" encoding="utf-8"?>\n')
                        f.write('<tt xmlns="http://www.w3.org/ns/ttml">\n<body>\n<div>\n')
                        for seg in segments:
                            def ttml_time(t):
                                h = int(t // 3600)
                                m = int((t % 3600) // 60)
                                s = int(t % 60)
                                ms = int((t - int(t)) * 1000)
                                return f"{h:02}:{m:02}:{s:02}.{ms:03}"
                            f.write(f'<p begin="{ttml_time(seg["start"])}" end="{ttml_time(seg["end"])}">{seg["text"].strip()}</p>\n')
                        f.write('</div>\n</body>\n</tt>\n')
                    print(f"Saved DFXP/TTML to {out_path}")
                elif job == 'cue':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.cue")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(f'FILE "{base}.mp3" MP3\n')
                        for idx, seg in enumerate(segments, 1):
                            m, s = divmod(int(seg['start']), 60)
                            f.write(f'  TRACK {idx:02d} AUDIO\n    INDEX 01 {m:02d}:{s:02d}:00\n')
                    print(f"Saved CUE to {out_path}")
                elif job == 'chapters':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.chapters.txt")
                    segments = result.get('segments', [{'start': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        for idx, seg in enumerate(segments, 1):
                            m, s = divmod(int(seg['start']), 60)
                            f.write(f'CHAPTER{idx:02d}={m:02d}:{s:02d}.000\n')
                            f.write(f'CHAPTER{idx:02d}NAME={seg["text"].strip()}\n')
                    print(f"Saved chapters to {out_path}")
                elif job == 'jsonl':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.jsonl")
                    segments = result.get('segments', [{'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        for seg in segments:
                            json.dump(seg, f, ensure_ascii=False)
                            f.write('\n')
                    print(f"Saved JSONL to {out_path}")
                elif job == 'stats':
                    word_count = len(result['text'].split())
                    duration = sum(seg.get('end', 0) - seg.get('start', 0) for seg in result.get('segments', []))
                    stats = {
                        'file': file_path,
                        'word_count': word_count,
                        'duration': duration,
                        'segment_count': len(result.get('segments', [])),
                    }
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.stats.json")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        json.dump(stats, f, ensure_ascii=False, indent=2)
                    print(f"Saved stats JSON to {out_path}")
                elif job == 'frequency':
                    import csv
                    import collections
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.freq.csv")
                    words = result['text'].lower().split()
                    freq = collections.Counter(words)
                    with open(out_path, 'w', encoding='utf-8', newline='') as f:
                        writer = csv.writer(f)
                        writer.writerow(['Word', 'Frequency'])
                        for word, count in freq.most_common():
                            writer.writerow([word, count])
                    print(f"Saved frequency CSV to {out_path}")
                elif job == 'rtf':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.rtf")
                    # Simple RTF: wrap text in RTF header/footer
                    rtf_text = result['text'].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                    rtf_content = f"{{\\rtf1\\ansi\\deff0{{\\fonttbl{{\\f0 Courier;}}}}\\f0\\fs20 {rtf_text}}}"
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(rtf_content)
                    print(f"Saved RTF to {out_path}")
                elif job == 'xlsx':
                    try:
                        import openpyxl
                    except ImportError:
                        print("openpyxl not installed. Please install with 'pip install openpyxl'.")
                        continue
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.xlsx")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    wb = openpyxl.Workbook()
                    ws = wb.active
                    ws.title = "Transcript"
                    ws.append(['Index', 'Start', 'End', 'Text'])
                    for idx, seg in enumerate(segments, 1):
                        ws.append([idx, seg.get('start', ''), seg.get('end', ''), seg.get('text', '')])
                    wb.save(out_path)
                    print(f"Saved XLSX to {out_path}")
                elif job == 'json_srt':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.json.srt")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    srt_json = []
                    for idx, seg in enumerate(segments, 1):
                        srt_json.append({
                            'index': idx,
                            'start': seg.get('start', 0),
                            'end': seg.get('end', 0),
                            'text': seg.get('text', '')
                    })
                    with open(out_path, 'w', encoding='utf-8') as f:
                        json.dump(srt_json, f, ensure_ascii=False, indent=2)
                    print(f"Saved JSON SRT to {out_path}")
                elif job == 'vxml':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.vxml")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    vxml = [
                        '<?xml version="1.0" encoding="UTF-8"?>',
                        '<vxml version="2.1">',
                        '  <form id="transcript">'
                    ]
                    for idx, seg in enumerate(segments, 1):
                        vxml.append(f'    <block>{seg.get("text", "").strip()}</block>')
                    vxml.append('  </form>')
                    vxml.append('</vxml>')
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write('\n'.join(vxml))
                    print(f"Saved VoiceXML to {out_path}")
                elif job == 'edi':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.edi")
                    # Very simple EDI: one segment per line, pipe-separated
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    with open(out_path, 'w', encoding='utf-8') as f:
                        for idx, seg in enumerate(segments, 1):
                            f.write(f"SEG|{idx}|{seg.get('start', 0)}|{seg.get('end', 0)}|{seg.get('text', '').replace('|', ' ')}\n")
                    print(f"Saved EDI to {out_path}")
                elif job == 'audio_aligned_json':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.audio_aligned.json")
                    # Use word-level timestamps if available, else segment-level
                    words = result.get('words')
                    if not words:
                        # Fallback: use segments
                        words = []
                        for seg in result.get('segments', []):
                            words.append({
                                'word': seg.get('text', ''),
                                'start': seg.get('start', 0),
                                'end': seg.get('end', 0),
                                'speaker': seg.get('speaker', '')
                            })
                    with open(out_path, 'w', encoding='utf-8') as f:
                        json.dump(words, f, ensure_ascii=False, indent=2)
                    print(f"Saved audio-aligned JSON to {out_path}")
                elif job == 'eaf':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.eaf")
                    # Minimal ELAN .eaf XML structure
                    import xml.etree.ElementTree as ET
                    root = ET.Element('ANNOTATION_DOCUMENT', AUTHOR="Whisper", DATE="", FORMAT="3.0", VERSION="2.8")
                    time_order = ET.SubElement(root, 'TIME_ORDER')
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    for idx, seg in enumerate(segments, 1):
                        ET.SubElement(time_order, 'TIME_SLOT', TIME_SLOT_ID=f"ts{2*idx-1}", TIME_VALUE=str(int(seg['start']*1000)))
                        ET.SubElement(time_order, 'TIME_SLOT', TIME_SLOT_ID=f"ts{2*idx}", TIME_VALUE=str(int(seg['end']*1000)))
                    tier = ET.SubElement(root, 'TIER', TIER_ID="transcript", LINGUISTIC_TYPE_REF="default-lt")
                    for idx, seg in enumerate(segments, 1):
                        ann = ET.SubElement(tier, 'ANNOTATION')
                        align = ET.SubElement(ann, 'ALIGNABLE_ANNOTATION', ANNOTATION_ID=f"a{idx}", TIME_SLOT_REF1=f"ts{2*idx-1}", TIME_SLOT_REF2=f"ts{2*idx}")
                        ET.SubElement(align, 'ANNOTATION_VALUE').text = seg['text']
                    ET.SubElement(root, 'LINGUISTIC_TYPE', LINGUISTIC_TYPE_ID="default-lt", TIME_ALIGNABLE="true")
                    tree = ET.ElementTree(root)
                    tree.write(out_path, encoding='utf-8', xml_declaration=True)
                    print(f"Saved ELAN EAF to {out_path}")
                elif job == 'textgrid':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}.TextGrid")
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    # Praat TextGrid format (simple, one tier)
                    xmin = 0.0
                    xmax = max(seg.get('end', 0) for seg in segments) if segments else 0.0
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(f"File type = \"ooTextFile\"\nObject class = \"TextGrid\"\n\n")
                        f.write(f"xmin = {xmin}\nxmax = {xmax}\ntiers? <exists>\nsize = 1\nitem []:\n    item [1]:\n        class = \"IntervalTier\"\n        name = \"transcript\"\n        xmin = {xmin}\n        xmax = {xmax}\n        intervals: size = {len(segments)}\n")
                        for idx, seg in enumerate(segments, 1):
                            f.write(f"        intervals [{idx}]:\n            xmin = {seg.get('start', 0)}\n            xmax = {seg.get('end', 0)}\n            text = \"{seg.get('text', '').replace('\\', ' ').replace('"', '')}\"\n")
                    print(f"Saved Praat TextGrid to {out_path}")
                elif job == 'webembed':
                    out_path = os.path.join(out_dir, f"{base}{lang_suffix}_webembed.html")
                    # Simple HTML+JS: audio player + transcript panel
                    audio_file = os.path.basename(file_path)
                    segments = result.get('segments', [{'start': 0, 'end': 0, 'text': result['text']}])
                    transcript_html = ''
                    for seg in segments:
                        start = seg.get('start', 0)
                        transcript_html += f'<span class="transcript-seg" data-start="{start}">{seg.get("text", "").strip()}</span> '
                    html = f"""
                        <html><head><meta charset='utf-8'><title>{base} WebEmbed</title>
                        <style>.transcript-seg{{cursor:pointer;}}.transcript-seg:hover{{background:yellow;}}</style>
                        </head><body>
                        <audio id='audio' controls src='{audio_file}'></audio>
                        <div id='transcript'>{transcript_html}</div>
                        <script>
                        document.querySelectorAll('.transcript-seg').forEach(function(span){{
                          span.onclick = function(){{
                            var audio = document.getElementById('audio');
                            audio.currentTime = parseFloat(span.getAttribute('data-start'));
                            audio.play();
                          }};
                        }});
                        </script>
                        </body></html>
                    """
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(html)
                    print(f"Saved WebEmbed HTML to {out_path}")
                else:
                    print(f"[DEBUG] Unknown job type encountered: '{job}' for file: {file_path}")
        # Always output original text
        output_jobs(jobs, base, result['text'], segments)
        # Only output translations if translations is a non-empty dict
        if isinstance(translations, dict) and translations:
            for lang, translated_text in translations.items():
                print(f"[DEBUG] Translation output: lang={lang!r}, translated_text={'present' if translated_text else 'empty or None'}")
                if not lang or lang is None or str(lang).lower() == "none":
                    continue
                if not translated_text or translated_text is None:
                    continue
                lang_suffix = f"_{lang}" if len(translations) > 1 else f"_{lang}"
                # For SRT, use per-segment translations
                if segment_translations and lang in segment_translations:
                    output_jobs(jobs, base, translated_text, segments, lang_suffix, translated_segments=segment_translations[lang])
                else:
                    output_jobs(jobs, base, translated_text, None, lang_suffix)
        results.append({'file': file_path, 'text': result['text']})
    _tasks.clear()
    print("All tasks completed.")
    return results

if __name__ == "__main__":
    model = get_whisper_model()
    # Example usage: transcribe an audio file
    # result = model.transcribe("your_audio_file.wav")
    # print(result["text"])
